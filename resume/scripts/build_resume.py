from __future__ import annotations

import argparse
import json
import re
import shutil
import subprocess
import sys
import tempfile
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


SCRIPT_DIR = Path(__file__).resolve().parent
PROJECT_DIR = SCRIPT_DIR.parent
REPO_ROOT = PROJECT_DIR.parent
VARIANTS_DIR = PROJECT_DIR / "variants"
FULL_SOURCE = VARIANTS_DIR / "full.md"
APPLICATION_SOURCE = VARIANTS_DIR / "application.md"
DOCS_DIR = REPO_ROOT / "docs"
SITE_DIR = REPO_ROOT / "site"
DOWNLOADS_DIR = DOCS_DIR / "downloads"
RESUME_INCLUDE_OUTPUT = SITE_DIR / "_includes" / "generated" / "resume-body.html"
RESUME_DATA_OUTPUT = SITE_DIR / "_data" / "resume.json"
FULL_DOCX_OUTPUT = DOWNLOADS_DIR / "bt-franklin-resume-full.docx"
FULL_PDF_OUTPUT = DOWNLOADS_DIR / "bt-franklin-resume-full.pdf"
APPLICATION_DOCX_OUTPUT = DOWNLOADS_DIR / "bt-franklin-resume-application.docx"
APPLICATION_PDF_OUTPUT = DOWNLOADS_DIR / "bt-franklin-resume-application.pdf"
BODY_FONT = "Helvetica Neue"
BODY_COLOR = RGBColor(0x0C, 0x0C, 0x0C)
LINK_COLOR = RGBColor(0x05, 0x63, 0xC1)
SECTION_SIZES = {
    "Core Expertise": 12,
    "Additional Experience": 13,
}


class BuildError(RuntimeError):
    pass


def require_tool(name: str) -> str:
    tool = shutil.which(name)
    if tool is None:
        raise BuildError(f"Required command not found: {name}")
    return tool


def run(command: list[str], *, cwd: Path | None = None, capture: bool = False) -> str:
    result = subprocess.run(
        command,
        cwd=cwd,
        check=False,
        stdout=subprocess.PIPE if capture else None,
        stderr=subprocess.PIPE,
        text=True,
    )
    if result.returncode != 0:
        stderr = result.stderr.strip()
        raise BuildError(f"Command failed: {' '.join(command)}\n{stderr}")
    return result.stdout if capture else ""


def parse_frontmatter(source: str) -> tuple[dict[str, str], str]:
    if not source.startswith("---\n"):
        return {}, source

    end = source.find("\n---\n", 4)
    if end == -1:
        raise BuildError("Opening frontmatter marker found without closing marker.")

    frontmatter = source[4:end]
    body = source[end + len("\n---\n") :]
    metadata: dict[str, str] = {}
    for line in frontmatter.splitlines():
        if not line.strip():
            continue
        key, separator, value = line.partition(":")
        if not separator:
            raise BuildError(f"Unsupported frontmatter line: {line}")
        metadata[key.strip()] = value.strip().strip('"').strip("'")
    return metadata, body


def set_run_font(run, *, size: int = 12, bold: bool | None = None, italic: bool | None = None) -> None:
    run.font.name = BODY_FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
    run.font.size = Pt(size)
    run.font.color.rgb = BODY_COLOR
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_paragraph_spacing(paragraph, *, before: int = 0, after: int = 8) -> None:
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)


def add_hyperlink(paragraph, text: str, url: str) -> None:
    part = paragraph.part
    relationship_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)

    run = OxmlElement("w:r")
    run_properties = OxmlElement("w:rPr")

    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    run_properties.append(color)

    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), BODY_FONT)
    fonts.set(qn("w:hAnsi"), BODY_FONT)
    fonts.set(qn("w:eastAsia"), BODY_FONT)
    run_properties.append(fonts)

    size = OxmlElement("w:sz")
    size.set(qn("w:val"), "24")
    run_properties.append(size)

    text_element = OxmlElement("w:t")
    text_element.text = text
    run.append(run_properties)
    run.append(text_element)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def clean_inline(text: str) -> str:
    return text.replace("\\", "").strip()


def add_markdown_runs(paragraph, text: str, *, default_italic: bool = False) -> None:
    text = clean_inline(text)
    pattern = re.compile(r"\[([^\]]+)\]\(([^)]+)\)|\*\*([^*]+)\*\*|\*([^*]+)\*|<([^>]+)>")
    position = 0
    for match in pattern.finditer(text):
        if match.start() > position:
            run = paragraph.add_run(text[position : match.start()])
            set_run_font(run, italic=default_italic)
        if match.group(1) is not None:
            add_hyperlink(paragraph, match.group(1), match.group(2))
        elif match.group(3) is not None:
            run = paragraph.add_run(match.group(3))
            set_run_font(run, bold=True, italic=default_italic)
        elif match.group(4) is not None:
            run = paragraph.add_run(match.group(4))
            set_run_font(run, italic=True)
        elif match.group(5) is not None:
            value = match.group(5)
            url = f"mailto:{value}" if "@" in value and not value.startswith("mailto:") else value
            add_hyperlink(paragraph, value.removeprefix("mailto:"), url)
        position = match.end()
    if position < len(text):
        run = paragraph.add_run(text[position:])
        set_run_font(run, italic=default_italic)


def parse_markdown_blocks(body: str) -> list[tuple[str, int | None, str]]:
    lines = body.splitlines()
    blocks: list[tuple[str, int | None, str]] = []
    index = 0
    while index < len(lines):
        line = lines[index]
        if not line.strip():
            index += 1
            continue

        heading = re.match(r"^(#{1,6})\s+(.+)$", line)
        if heading:
            blocks.append(("heading", len(heading.group(1)), heading.group(2).strip()))
            index += 1
            continue

        if line.startswith("- "):
            item = line[2:].strip()
            index += 1
            while index < len(lines) and lines[index].startswith("  "):
                item += " " + lines[index].strip()
                index += 1
            blocks.append(("bullet", None, item))
            continue

        parts = [line.strip()]
        index += 1
        while index < len(lines):
            next_line = lines[index]
            if not next_line.strip() or next_line.startswith("- ") or re.match(r"^(#{1,6})\s+", next_line):
                break
            parts.append(next_line.strip())
            index += 1
        blocks.append(("paragraph", None, " ".join(parts)))
    return blocks


def add_normal_paragraph(document: Document, text: str, *, before: int = 0, after: int = 8) -> None:
    paragraph = document.add_paragraph()
    set_paragraph_spacing(paragraph, before=before, after=after)
    add_markdown_runs(paragraph, text)


def add_bullet(document: Document, text: str, *, before: int = 0, after: int = 8) -> None:
    paragraph = document.add_paragraph(style="List Bullet")
    set_paragraph_spacing(paragraph, before=before, after=after)
    add_markdown_runs(paragraph, text)


def add_section_heading(document: Document, text: str, *, before: int = 16) -> None:
    paragraph = document.add_paragraph()
    set_paragraph_spacing(paragraph, before=before, after=8)
    run = paragraph.add_run(text)
    set_run_font(run, size=SECTION_SIZES.get(text, 16), bold=True)


def add_bold_heading(document: Document, text: str, *, before: int = 0, after: int = 8) -> None:
    paragraph = document.add_paragraph()
    set_paragraph_spacing(paragraph, before=before, after=after)
    run = paragraph.add_run(text)
    set_run_font(run, bold=True)


def add_role_paragraph(document: Document, title: str, dates: str) -> None:
    paragraph = document.add_paragraph()
    set_paragraph_spacing(paragraph)
    title_run = paragraph.add_run(title)
    set_run_font(title_run, italic=True)
    paragraph.add_run().add_break()
    date_run = paragraph.add_run(dates)
    set_run_font(date_run, italic=True)


def add_product_heading(document: Document, name: str, detail: str) -> None:
    paragraph = document.add_paragraph()
    set_paragraph_spacing(paragraph)
    name_run = paragraph.add_run(name)
    set_run_font(name_run, bold=True)
    separator = paragraph.add_run(" - ")
    set_run_font(separator)
    detail_run = paragraph.add_run(detail)
    set_run_font(detail_run, italic=True)


def add_education_heading(document: Document, school: str, degree: str, minor: str) -> None:
    paragraph = document.add_paragraph()
    set_paragraph_spacing(paragraph)
    school_run = paragraph.add_run(school)
    set_run_font(school_run, bold=True)
    paragraph.add_run().add_break()
    degree_run = paragraph.add_run(degree)
    set_run_font(degree_run)
    paragraph.add_run().add_break()
    minor_run = paragraph.add_run(minor)
    set_run_font(minor_run)


def add_contact_header(document: Document, blocks: list[tuple[str, int | None, str]]) -> int:
    name = blocks[0][2]
    city = blocks[1][2]
    headline = blocks[2][2]
    contact = blocks[3][2]
    links = blocks[4][2]

    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(paragraph, after=24)

    run = paragraph.add_run(name)
    set_run_font(run, bold=True)
    paragraph.add_run().add_break()

    run = paragraph.add_run(city)
    set_run_font(run)
    paragraph.add_run().add_break()

    run = paragraph.add_run(headline)
    set_run_font(run)
    paragraph.add_run().add_break()

    phone, email_text = [part.strip() for part in contact.split("|", 1)]
    run = paragraph.add_run(f"{phone} | ")
    set_run_font(run)
    add_markdown_runs(paragraph, email_text)
    paragraph.add_run().add_break()

    link_matches = re.findall(r"\[([^\]]+)\]\(([^)]+)\)", links)
    for idx, (label, url) in enumerate(link_matches):
        if idx:
            run = paragraph.add_run(" \u2022 ")
            set_run_font(run)
        add_hyperlink(paragraph, label, url)

    return 5


def build_styled_docx(body: str, output_path: Path) -> None:
    blocks = parse_markdown_blocks(body)
    if len(blocks) < 5:
        raise BuildError("Resume source is missing the expected header blocks.")

    document = Document()
    section = document.sections[0]
    for margin in ["top_margin", "right_margin", "bottom_margin", "left_margin"]:
        setattr(section, margin, Inches(0.7875))

    normal_style = document.styles["Normal"]
    normal_style.font.name = BODY_FONT
    normal_style._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
    normal_style.font.size = Pt(12)
    normal_style.font.color.rgb = BODY_COLOR

    index = add_contact_header(document, blocks)
    current_section = ""
    while index < len(blocks):
        kind, level, text = blocks[index]
        if kind == "heading" and level == 2:
            current_section = text
            add_section_heading(document, text, before=0 if text == "Professional Summary" else 16)
            index += 1
            continue
        if kind == "heading" and level == 3 and current_section == "Selected Founder-Led AI Products":
            detail = blocks[index + 1][2] if index + 1 < len(blocks) else ""
            add_product_heading(document, text, detail)
            index += 2
            continue
        if kind == "heading" and level == 3 and current_section == "Education":
            degree = blocks[index + 1][2] if index + 1 < len(blocks) else ""
            minor = blocks[index + 2][2] if index + 2 < len(blocks) else ""
            add_education_heading(document, text, degree, minor)
            index += 3
            continue
        if kind == "heading" and level == 3:
            add_bold_heading(document, text)
            index += 1
            continue
        if kind == "heading" and level == 4:
            dates = blocks[index + 1][2] if index + 1 < len(blocks) else ""
            add_role_paragraph(document, text, dates)
            index += 2
            continue
        if kind == "bullet":
            add_bullet(document, text)
            index += 1
            continue
        if kind == "paragraph":
            add_normal_paragraph(document, text)
            index += 1
            continue
        index += 1

    document.save(output_path)


def convert_markdown_to_html(source_path: Path) -> tuple[dict[str, str], str]:
    source = source_path.read_text(encoding="utf-8")
    metadata, _body = parse_frontmatter(source)
    body_html = run(
        [
            "pandoc",
            str(source_path),
            "--from=gfm+yaml_metadata_block",
            "--to=html",
        ],
        capture=True,
    ).strip()
    return metadata, body_html


def build_resume_site_artifacts(metadata: dict[str, str], body_html: str) -> None:
    RESUME_INCLUDE_OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    RESUME_DATA_OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    RESUME_INCLUDE_OUTPUT.write_text(body_html + "\n", encoding="utf-8")
    RESUME_DATA_OUTPUT.write_text(
        json.dumps(
            {
                "title": metadata.get("title", "B.T. Franklin Resume"),
                "last_updated": metadata.get("last_updated", ""),
            },
            indent=2,
        )
        + "\n",
        encoding="utf-8",
    )


def build_docx(source_path: Path, output_path: Path) -> None:
    if output_path.exists() and output_path.stat().st_mtime >= source_path.stat().st_mtime:
        return
    _metadata, body = parse_frontmatter(source_path.read_text(encoding="utf-8"))
    build_styled_docx(body, output_path)


def build_pdf(docx_output: Path) -> None:
    pdf_output = docx_output.with_suffix(".pdf")
    if pdf_output.exists() and pdf_output.stat().st_mtime >= docx_output.stat().st_mtime:
        return
    with tempfile.TemporaryDirectory(prefix="btfranklin-lo-") as profile_dir:
        run(
            [
                "soffice",
                f"-env:UserInstallation=file://{profile_dir}",
                "--headless",
                "--convert-to",
                "pdf",
                "--outdir",
                str(DOWNLOADS_DIR),
                str(docx_output),
            ]
        )


def validate_outputs() -> None:
    for output in [
        RESUME_INCLUDE_OUTPUT,
        RESUME_DATA_OUTPUT,
        FULL_DOCX_OUTPUT,
        FULL_PDF_OUTPUT,
        APPLICATION_DOCX_OUTPUT,
        APPLICATION_PDF_OUTPUT,
    ]:
        if not output.exists():
            raise BuildError(f"Expected output was not created: {output}")
        if output.stat().st_size == 0:
            raise BuildError(f"Expected output is empty: {output}")


def build() -> None:
    require_tool("pandoc")
    require_tool("soffice")
    for source_path in [FULL_SOURCE, APPLICATION_SOURCE]:
        if not source_path.exists():
            raise BuildError(f"Resume source not found: {source_path}")

    DOWNLOADS_DIR.mkdir(parents=True, exist_ok=True)

    metadata, body_html = convert_markdown_to_html(FULL_SOURCE)
    build_resume_site_artifacts(metadata, body_html)
    build_docx(FULL_SOURCE, FULL_DOCX_OUTPUT)
    build_docx(APPLICATION_SOURCE, APPLICATION_DOCX_OUTPUT)
    build_pdf(FULL_DOCX_OUTPUT)
    build_pdf(APPLICATION_DOCX_OUTPUT)
    validate_outputs()


def main() -> int:
    parser = argparse.ArgumentParser(description="Build resume web, DOCX, and PDF outputs.")
    parser.add_argument("--check", action="store_true", help="Build and verify expected outputs.")
    parser.parse_args()

    try:
        build()
    except BuildError as error:
        print(error, file=sys.stderr)
        return 1

    print(f"Wrote {RESUME_INCLUDE_OUTPUT.relative_to(REPO_ROOT)}")
    print(f"Wrote {RESUME_DATA_OUTPUT.relative_to(REPO_ROOT)}")
    print(f"Wrote {FULL_DOCX_OUTPUT.relative_to(REPO_ROOT)}")
    print(f"Wrote {FULL_PDF_OUTPUT.relative_to(REPO_ROOT)}")
    print(f"Wrote {APPLICATION_DOCX_OUTPUT.relative_to(REPO_ROOT)}")
    print(f"Wrote {APPLICATION_PDF_OUTPUT.relative_to(REPO_ROOT)}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
