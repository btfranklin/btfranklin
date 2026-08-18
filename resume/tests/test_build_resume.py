from __future__ import annotations

import tempfile
import unittest
from pathlib import Path

from docx import Document

from scripts import build_resume


HEADER = """# B.T. Franklin

Phoenix, Arizona

602-247-0439 | <brandon.franklin@gmail.com>

[btfranklin.info](https://btfranklin.info/)
"""


class ResumeBuilderTests(unittest.TestCase):
    def build_document(self, body: str, *, headline: str | None = None) -> Document:
        with tempfile.TemporaryDirectory() as temp_dir:
            output_path = Path(temp_dir) / "resume.docx"
            build_resume.build_styled_docx(body, output_path, headline=headline)
            return Document(output_path)

    def test_header_without_headline_keeps_first_section(self) -> None:
        document = self.build_document(f"{HEADER}\n## Professional Summary\n\nShort summary.\n")
        text = "\n".join(paragraph.text for paragraph in document.paragraphs)
        self.assertIn("B.T. Franklin\nPhoenix, Arizona\n602-247-0439", text)
        self.assertIn("Professional Summary", text)
        self.assertNotIn("Applied AI Technical Lead", text)

    def test_optional_headline_is_rendered(self) -> None:
        document = self.build_document(
            f"{HEADER}\n## Professional Summary\n\nShort summary.\n",
            headline="Applied AI Technical Lead",
        )
        self.assertIn("Applied AI Technical Lead", document.paragraphs[0].text)

    def test_degree_only_education_preserves_next_section(self) -> None:
        body = f"""{HEADER}

## Education

### Transylvania University

BA in Computer Science, 1997

## Skills

Python
"""
        document = self.build_document(body)
        text = "\n".join(paragraph.text for paragraph in document.paragraphs)
        self.assertIn("Transylvania University\nBA in Computer Science, 1997", text)
        self.assertIn("Skills", text)
        self.assertIn("Python", text)

    def test_optional_education_detail_is_rendered(self) -> None:
        body = f"""{HEADER}

## Education

### Transylvania University

BA in Computer Science, 1997

Minor in Psychology
"""
        document = self.build_document(body)
        self.assertIn("Minor in Psychology", document.paragraphs[-1].text)

    def test_malformed_header_fails_closed(self) -> None:
        body = """# B.T. Franklin

Phoenix, Arizona

602-247-0439 | <brandon.franklin@gmail.com>

## Professional Summary

Short summary.
"""
        with tempfile.TemporaryDirectory() as temp_dir:
            with self.assertRaises(build_resume.BuildError):
                build_resume.build_styled_docx(body, Path(temp_dir) / "resume.docx")

    def test_long_bullet_reports_source_line(self) -> None:
        source = "- " + "word " * 36
        findings = build_resume.find_long_bullets(source)
        self.assertEqual([(1, 36)], [(line, count) for line, count, _item in findings])


if __name__ == "__main__":
    unittest.main()
